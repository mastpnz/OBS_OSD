import sys
import re
from docx.api import Document

f = open(sys.argv[1], 'rb')
document = Document(f)

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
document = Document(f)
table = document.tables[0]

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []

#keys = None
#for i, row in enumerate(table.rows):
#    text = (cell.text for cell in row.cells)
#
#    # Establish the mapping based on the first row
#    # headers; these will become the keys of our dictionary
#    if i == 0:
#        keys = tuple(text)
#        continue
#
#    # Construct a dictionary for this row, mapping
#    # keys to values for this row
#    row_data = dict(zip(keys, text))
#    data.append(row_data)


fo = open('squad.json', 'w')
fo.write('{\n')
teamparse = re.search(r'^Заявочный лист (.+), (.+)', document.paragraphs[0].text)
fo.write('\t\"Name\" : \"')
fo.write(teamparse.group(1))
fo.write('\",\n')
fo.write('\t\"City\" : \"')
fo.write(teamparse.group(2))
fo.write('\",\n')
fo.write('\t\"Players\" : [\n')
for i, row in enumerate(table.rows):
    if i > 0:
        if i > 1:
            fo.write(',\n')
        str = '\t\t{ "number\" : \"' + row.cells[0].text + '\", ' + '\"FIO\" : \"' + row.cells[1].text + '\", ' + '\"amplua\" : \"' + row.cells[2].text + '\", ' + '\"bdate\" : \"' + row.cells[4].text + '\" }'  
        fo.write(str)
fo.write('\n')
fo.write('\t]')

table = document.tables[1]
fo.write(',\n')    
fo.write('\t\"Stuff\" : [\n')
for i, row in enumerate(table.rows):
    if i > 0:
        if i > 1:
            fo.write(',\n')
        str = '\t\t{ \"amplua\" : ' + row.cells[2].text + '\", ' + '\"FIO\" : \"' + row.cells[1].text + '\", ' + '\"telnumb\" : \"' + row.cells[3].text + '\", ' + '\"email\" : \"' + row.cells[4].text + '\" }'
        fo.write(str)
fo.write('\n')
fo.write('\t]\n')
fo.write('}\n')
fo.close()


#    print('строка №',i)
#    for j,cell in enumerate(row.cells):
#        print('ячейка №',j,'=',cell.text)

#print(*data)

f.close()
